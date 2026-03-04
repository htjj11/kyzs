from api.sql_role import kyzs_sql
import base64
from pdf2docx import parse
from docx import Document
from docx2pdf import convert
import json
import random

# 把上级目录加入搜索路径
# sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


from api.api_bg import siliconflow_deepseek_answer
from config import * 

'''
翻译词汇表
CREATE TABLE `translate_words` (
  `id` int(11) NOT NULL AUTO_INCREMENT,
  `ts_type` varchar(255) DEFAULT NULL COMMENT '该词汇用于的翻译类型，比如中英、中法等',
  `field_id` int(11) DEFAULT NULL COMMENT '绑定到translate_dict_field的主键，表明这个词汇是哪种领域',
  `content1` text COMMENT '单词语言1（通常为英文）',
  `content2` text COMMENT '单词释义2（通常为中文释义的简写）',
  `content3` text COMMENT '单词用法3（包括注释，以及什么情况下该用什么content2）',
  `from` text COMMENT '该条单词的来源（比如xxxx字典）',
  PRIMARY KEY (`id`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci;
'''

def translate_text_api(raw_text:str,translate_type:str,field_id:int):
    """
    翻译单个字符串的函数
    :param raw_text: 原始文本
    :param translate_type: 翻译类型,en2zh:英文到中文,zh2en:中文到英文
    :param field_id: 领域id
    :return: 翻译后的文本
    """
    print('用户请求翻译:',raw_text[:40],'...',)
    #定义用户传入的翻译类型与数据库的词典类型对应，比如中英或英中都是用en的词典
    word_ts_type = {
        'en2zh':'en',
        'zh2en':'en',
    }
    ts_type = word_ts_type[translate_type]
    sql = f"select * from translate_words where ts_type='{ts_type}' and field_id={field_id}"
    words_dict = kyzs_sql.mysql_exec(sql)
    new_words_dict = []

    #遍历数据库字典，判断raw_text是否有这个单词，如果有，在新字典中增加此记录。
    #此处需要判断是zh2en还是en2zh，因为zh2en的content1是中文，content2是英文，定义一个字典，值为索引
    word_index =  {'en2zh':'content1','zh2en':'content2'}
    import re
    # 将段落转为小写，便于匹配
    lower_paragraph = raw_text.lower()
    for i in words_dict:
        #判断单词是否在原始文本中，需要在原始文本中添加空格，防止匹配到其他单词的子字符串
        lower_word = i[word_index[translate_type]].lower()
        pattern = r'\b' + re.escape(lower_word) + r'\b'
        if re.search(pattern, lower_paragraph):
            new_words_dict.append(i)
            
    print('知识库命中字典个数:',new_words_dict.__len__())
    #将剩余字典的内容按照格式组成字符串，作为大模型提示词
    #格式 ：单词1：abandon 中文释义1：放弃，解释：放弃某些事情 单词2：abandon 中文释义2：放弃
    prompt = '以下是文中存在的词汇对照组，可用于参考：\n'
    for i in new_words_dict:
        prompt += f"{i[word_index[translate_type]]}：{i['content2']}，解释：{i['content3']}  "
    prompt += f"\n请翻译以下文本（单词）为中文：【{raw_text}】，输出格式为json：origin_text为原文,translate_text为译文，不要有任何其他信息、提问，只返回json即可"

    #循环遍历，直到大模型返回的json格式正确
    while 1:
        res = siliconflow_deepseek_answer(prompt)
        try:
            #大模型返回的结果样式：```json
            # {
            # "origin_text": "Abstract",
            # "translate_text": "摘要"
            # }
            # ```
            res = res.replace('```json','').replace('```','')
            res_json = json.loads(res)
            translate_text = res_json['translate_text']
            print('大模型翻译成功：',translate_text)
            break
        except Exception as e:
            print('大模型返回的结果格式错误:错误信息:',e)
            continue

    return {'translate_result':translate_text,'words_dict':new_words_dict}

import concurrent.futures

def translate_text_list_api(raw_text_list:list,translate_type:str,field_id:int):
    """
    翻译函数
    :param raw_text_list: 原始文本列表
    :param translate_type: 翻译类型,en2zh:英文到中文,zh2en:中文到英文
    :param field_id: 领域id
    :return: 翻译后的文本列表
    """
    translate_results = [None] * len(raw_text_list)
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # 使用多线程调用翻译接口
        futures = {executor.submit(translate_text_api, text, translate_type, field_id): index for index, text in enumerate(raw_text_list)}
        for future in concurrent.futures.as_completed(futures):
            index = futures[future]
            try:
                result = future.result()
                translate_results[index] = result['translate_result']
            except Exception as e:
                print(f"翻译出错: {e}")
                translate_results[index] = "该段翻译失败"
    return translate_results

def create_new_translate_doc_mission(base64_pdf_string,topic,user_id):
    '''
    创建新的翻译任务，先创建翻译任务，然后执行翻译方法，翻译后保存到数据库中
    :param base64_pdf_string: pdf文件base64字符串
    :param topic: 翻译任务的主题
    :param user_id: 用户ID
    :return: ok
    数据库格式：
CREATE TABLE `translate_doc` (
  `id` int NOT NULL AUTO_INCREMENT,
  `status` int DEFAULT NULL COMMENT '翻译状态，0为未完成，1为已完成',
  `raw_base64` longtext CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci COMMENT '原始文档的base64',
  `output_pdf_base64` longtext CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci COMMENT '输出为pdf的base64',
  `output_docx_base64` longtext CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci COMMENT '输出为docx的base64',
  `name` text CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci COMMENT '任务的名称',
  `user_id` int DEFAULT NULL COMMENT '绑定的user_id',
  PRIMARY KEY (`id`) USING BTREE
) ENGINE=InnoDB AUTO_INCREMENT=86 DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_0900_ai_ci ROW_FORMAT=DYNAMIC;
    '''


    #创建新的翻译任务
    sql = f"insert into translate_doc (name,status,raw_base64,user_id) values ('{topic}',0,'{base64_pdf_string}','{user_id}')"

    kyzs_sql.mysql_exec(sql)
    #获取最新的任务id
    sql = "select id from translate_doc order by id desc limit 1"
    mission_id = kyzs_sql.mysql_exec(sql)[0]['id']
    #执行翻译方法，获取两个base64字符串
    output_docx_base64,output_pdf_base64 = translate_pdf(base64_pdf_string)
    #将翻译后的word，pdf分别存储在根目录下/file文件夹中，文件名命名为当前时间
    import time
    file_name = time.strftime("%Y%m%d%H%M%S", time.localtime())
    with open(f'./file/{file_name}.pdf', 'wb') as f:
        f.write(base64.b64decode(output_pdf_base64))
    with open(f'./file/{file_name}.docx', 'wb') as f:
        f.write(base64.b64decode(output_docx_base64))
    #在sql中更新任务状态，pdf、docx的文件相对路径+文件名称，
    sql = f"update translate_doc set status=1,output_pdf_base64='./file/{file_name}.pdf',output_docx_base64='./file/{file_name}.docx' where id={mission_id}"
    kyzs_sql.mysql_exec(sql)
    print('翻译任务id成功:',mission_id)

    return 'ok'



def translate_pdf(base64_pdf_string):
    # pdf2docx == 0.5.6
    import base64
    from pdf2docx import parse
    from docx import Document
    from docx2pdf import convert
    import os
    # Decode the Base64 PDF string
    try:
        pdf_binary = base64.b64decode(base64_pdf_string)
    except:
        return "Invalid Base64 string"
    
    # Create a temporary PDF file to hold the decoded data
    with open("temp.pdf", "wb") as f:
        f.write(pdf_binary)
    try:
        parse("temp.pdf","input.docx")
    except Exception as e :
        print(f"{e}")

    document = Document('input.docx')
    list_paragraphs = [] 
    # paragraph first iteration
    for paragraph in document.paragraphs:

        if any("pic:pic" in run.element.xml for run in paragraph.runs):
            continue
        init = 1
        for run in paragraph.runs:
            if init:
                init = 0
                list_paragraphs.append(paragraph.text)
            else:
                run.text = ""
    # sudo interface
    translated_list_paragraphs = translate_text_list_api(list_paragraphs,'en2zh',1)
    # paragraph second iteration
    p_index = 0
    for paragraph in document.paragraphs:
        if any("pic:pic" in run.element.xml for run in paragraph.runs):
            continue
        init = 1
        for run in paragraph.runs:
            if init:
                init = 0
                
                run.text = translated_list_paragraphs[p_index]
                p_index += 1
            else:
                run.text = ""
    document.save('output.docx')
    try:
        convert("output.docx", "output.pdf")
    except:
        pass
    with open("output.docx", "rb") as f:
        output_binary_docx = f.read()
    base64_output_string_docx = base64.b64encode(output_binary_docx).decode('utf-8')
    
    # convert_docx_to_pdf_libreoffice("output.docx", "output.pdf")
    
    with open("output.pdf", "rb") as f:
        output_binary_pdf = f.read()
    base64_output_string_pdf = base64.b64encode(output_binary_pdf).decode('utf-8')
    
    os.remove("temp.pdf")
    os.remove("input.docx")
    # os.remove("output.docx")
    # os.remove("output.pdf")
    
    return [base64_output_string_docx, base64_output_string_pdf]

def get_all_translate_doc_list_api(user_id):
    sql = f"select id,name,status from translate_doc where user_id={user_id}"
    res = kyzs_sql.mysql_exec(sql)
    return {'translate_doc_list':res}

def get_translate_doc_detail_api(doc_id):
    sql = f"select output_pdf_base64,output_docx_base64 from translate_doc where id={doc_id}"
    res = kyzs_sql.mysql_exec(sql)[0]

    #此处字典内的base64实际上是相对文件路径，需要先打开文件，再将其转换为base64
    with open(res['output_pdf_base64'], 'rb') as f:
        output_binary_pdf = f.read()
    base64_output_string_pdf = base64.b64encode(output_binary_pdf).decode('utf-8')

    with open(res['output_docx_base64'], 'rb') as f:
        output_binary_docx = f.read()
    base64_output_string_docx = base64.b64encode(output_binary_docx).decode('utf-8')

    #测试将docx保存到本地看一下能否打开
    with open("test.docx", "wb") as f:
        f.write(base64.b64decode(base64_output_string_docx))

    res['output_pdf_base64'] = base64_output_string_pdf
    res['output_docx_base64'] = base64_output_string_docx
    return {'translate_doc_detail':res}

def del_translate_doc_api(doc_id):
    #先查询对应id的文档路径
    sql = f"select output_pdf_base64,output_docx_base64 from translate_doc where id={doc_id}"
    res = kyzs_sql.mysql_exec(sql)[0]
    #删除文件
    import os
    os.remove(res['output_pdf_base64'])
    os.remove(res['output_docx_base64'])

    #删除数据库中的记录
    sql = f"delete from translate_doc where id={doc_id}"
    kyzs_sql.mysql_exec(sql)


    return 'ok'



def get_translate_word_by_content_api(content1: str):
    """
    根据content1或content2获取全部翻译词汇
    :param content1: 关键词
    :return: 翻译词汇详情
    """
    # 直接拼写SQL语句，不使用参数化查询
    # 对输入进行简单转义，避免SQL注入
    escaped_keyword = content1.replace("'", "''")
    sql = f"SELECT * FROM translate_words WHERE content1 LIKE '%{escaped_keyword}%' OR content2 LIKE '%{escaped_keyword}%' limit 50"
    res = kyzs_sql.mysql_exec(sql)
    if res:
        return {'code': 200, 'msg': 'success', 'data': res}
    else:
        return {'code': 404, 'msg': '词汇不存在', 'data': None}

def add_translate_word_api(ts_type: str, field_id: int, content1: str, content2: str, content3: str, from_source: str):
    """
    添加新的翻译词汇
    :param ts_type: 翻译类型
    :param field_id: 领域ID
    :param content1: 单词语言1（通常为英文）
    :param content2: 单词释义2（通常为中文释义的简写）
    :param content3: 单词用法3（包括注释，以及什么情况下该用什么content2）
    :param from_source: 该条单词的来源（比如xxxx字典）
    :return: 操作结果
    """
    sql = f"""
    INSERT INTO translate_words (ts_type, field_id, content1, content2, content3, `from`) 
    VALUES ('{ts_type}', {field_id}, '{content1}', '{content2}', '{content3}', '{from_source}')
    """
    try:
        kyzs_sql.mysql_exec(sql)
        # 获取新插入记录的ID
        sql = "SELECT id FROM translate_words ORDER BY id DESC LIMIT 1"
        new_id = kyzs_sql.mysql_exec(sql)[0]['id']
        return {'code': 200, 'msg': 'success', 'data': {'id': new_id}}
    except Exception as e:
        return {'code': 500, 'msg': f'添加失败: {str(e)}', 'data': None}

def update_translate_word_api(word_id: int, ts_type: str = None, field_id: int = None, 
                             content1: str = None, content2: str = None, content3: str = None, from_source: str = None):
    """
    更新翻译词汇
    :param word_id: 词汇ID
    :param ts_type: 翻译类型
    :param field_id: 领域ID
    :param content1: 单词语言1（通常为英文）
    :param content2: 单词释义2（通常为中文释义的简写）
    :param content3: 单词用法3（包括注释，以及什么情况下该用什么content2）
    :param from_source: 该条单词的来源（比如xxxx字典）
    :return: 操作结果
    """
    # 构建更新语句，只更新非None的字段
    update_fields = []
    if ts_type is not None:
        update_fields.append(f"ts_type = '{ts_type}'")
    if field_id is not None:
        update_fields.append(f"field_id = {field_id}")
    if content1 is not None:
        update_fields.append(f"content1 = '{content1}'")
    if content2 is not None:
        update_fields.append(f"content2 = '{content2}'")
    if content3 is not None:
        update_fields.append(f"content3 = '{content3}'")
    if from_source is not None:
        update_fields.append(f"`from` = '{from_source}'")
    
    if not update_fields:
        return {'code': 400, 'msg': '没有提供要更新的字段', 'data': None}
    
    sql = f"UPDATE translate_words SET {', '.join(update_fields)} WHERE id = {word_id}"
    try:
        kyzs_sql.mysql_exec(sql)
        return {'code': 200, 'msg': 'success', 'data': None}
    except Exception as e:
        return {'code': 500, 'msg': f'更新失败: {str(e)}', 'data': None}

def delete_translate_word_api(word_id: int):
    """
    删除翻译词汇
    :param word_id: 词汇ID
    :return: 操作结果
    """
    sql = f"DELETE FROM translate_words WHERE id = {word_id}"
    try:
        kyzs_sql.mysql_exec(sql)
        return {'code': 200, 'msg': 'success', 'data': None}
    except Exception as e:
        return {'code': 500, 'msg': f'删除失败: {str(e)}', 'data': None}
 
if __name__ == '__main__':
    list1 = ['1','2','3']
    print(translate_text_list_api(list1,'en2zh',1))





