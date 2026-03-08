from app.core.database import db as kyzs_sql
import base64
import json
import re
import concurrent.futures
import time
import os

from app.services.literature_service import siliconflow_deepseek_answer


def translate_text_api(raw_text: str, translate_type: str, field_id: int):
    """
    翻译单个字符串，流程：
        1. 从 translate_words 词典表中查找原文里出现的专业术语
        2. 将匹配到的术语对照组拼入提示词，辅助大模型保持行业术语一致性
        3. 调用 DeepSeek 翻译，要求返回 JSON 格式 {origin_text, translate_text}
        4. 解析失败时重试，直到得到合法 JSON

    translate_type: 'en2zh' 英译中 / 'zh2en' 中译英
    field_id: 专业领域 id，用于过滤对应领域的词典
    """
    print('用户请求翻译:', raw_text[:40], '...')
    word_index = {'en2zh': 'content1', 'zh2en': 'content2'}
    ts_type_map = {'en2zh': 'en', 'zh2en': 'en'}
    ts_type = ts_type_map[translate_type]

    words_dict = kyzs_sql.mysql_exec(
        "SELECT * FROM translate_words WHERE ts_type=%s AND field_id=%s",
        (ts_type, field_id)
    )

    # 用正则全词匹配，避免"drill"匹配到"drilled"等子串
    lower_paragraph = raw_text.lower()
    matched = [
        w for w in words_dict
        if re.search(r'\b' + re.escape(w[word_index[translate_type]].lower()) + r'\b', lower_paragraph)
    ]
    print('知识库命中字典个数:', len(matched))

    prompt = '以下是文中存在的词汇对照组，可用于参考：\n'
    for w in matched:
        prompt += f"{w[word_index[translate_type]]}：{w['content2']}，解释：{w['content3']}  "
    prompt += f"\n请翻译以下文本（单词）为中文：【{raw_text}】，输出格式为json：origin_text为原文,translate_text为译文，不要有任何其他信息、提问，只返回json即可"

    while True:
        res = siliconflow_deepseek_answer(prompt)
        try:
            res_json = json.loads(res.replace('```json', '').replace('```', ''))
            translate_text = res_json['translate_text']
            print('大模型翻译成功：', translate_text)
            break
        except Exception as e:
            print('大模型返回的结果格式错误:错误信息:', e)

    return {'translate_result': translate_text, 'words_dict': matched}


def translate_text_list_api(raw_text_list: list, translate_type: str, field_id: int):
    """
    多段文本并发翻译，使用线程池保持段落顺序。
    用于 PDF 文档逐段翻译场景，段落数量可能达到数百条。
    """
    results = [None] * len(raw_text_list)
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {
            executor.submit(translate_text_api, text, translate_type, field_id): i
            for i, text in enumerate(raw_text_list)
        }
        for future in concurrent.futures.as_completed(futures):
            i = futures[future]
            try:
                results[i] = future.result()['translate_result']
            except Exception as e:
                print(f"翻译出错: {e}")
                results[i] = "该段翻译失败"
    return results


def create_new_translate_doc_mission(base64_pdf_string, topic, user_id):
    """
    创建翻译任务并执行翻译，由路由层通过 BackgroundTasks 异步调用，不阻塞请求响应。
    执行完成后将翻译产物写入 ./file/ 目录，并将文件路径回写数据库（status=1 表示完成）。
    """
    kyzs_sql.mysql_exec(
        "INSERT INTO translate_doc (name, status, raw_base64, user_id) VALUES (%s, 0, %s, %s)",
        (topic, base64_pdf_string, user_id)
    )
    mission = kyzs_sql.mysql_exec("SELECT id FROM translate_doc ORDER BY id DESC LIMIT 1")
    mission_id = mission[0]['id']

    output_docx_base64, output_pdf_base64 = translate_pdf(base64_pdf_string)
    file_name = time.strftime("%Y%m%d%H%M%S", time.localtime())

    # 翻译产物存本地 ./file/ 目录，数据库只存路径
    with open(f'./file/{file_name}.pdf', 'wb') as f:
        f.write(base64.b64decode(output_pdf_base64))
    with open(f'./file/{file_name}.docx', 'wb') as f:
        f.write(base64.b64decode(output_docx_base64))

    kyzs_sql.mysql_exec(
        "UPDATE translate_doc SET status=1, output_pdf_base64=%s, output_docx_base64=%s WHERE id=%s",
        (f'./file/{file_name}.pdf', f'./file/{file_name}.docx', mission_id)
    )
    print('翻译任务id成功:', mission_id)
    return 'ok'


def translate_pdf(base64_pdf_string):
    """
    PDF 翻译主流程：
        1. base64 解码 → 写临时文件 temp.pdf
        2. pdf2docx 解析版式 → input.docx（保留图片、表格布局）
        3. 提取所有纯文本段落，过滤含图片的段落
        4. 并发翻译所有段落
        5. 将译文写回 docx，保持原始排版
        6. docx2pdf 转换为 PDF（Windows 依赖 Word，Linux 依赖 LibreOffice）
    """
    from pdf2docx import parse
    from docx import Document
    from docx2pdf import convert

    pdf_binary = base64.b64decode(base64_pdf_string)
    with open("temp.pdf", "wb") as f:
        f.write(pdf_binary)
    try:
        parse("temp.pdf", "input.docx")
    except Exception as e:
        print(f"PDF解析出错: {e}")

    document = Document('input.docx')
    list_paragraphs = []
    for paragraph in document.paragraphs:
        # 含 pic:pic 标签的段落是图片，跳过不翻译
        if any("pic:pic" in run.element.xml for run in paragraph.runs):
            continue
        init = True
        for run in paragraph.runs:
            if init:
                init = False
                list_paragraphs.append(paragraph.text)
            else:
                # 同一段落多个 run 合并为第一个，其余清空，保持格式不乱
                run.text = ""

    translated = translate_text_list_api(list_paragraphs, 'en2zh', 1)
    p_index = 0
    for paragraph in document.paragraphs:
        if any("pic:pic" in run.element.xml for run in paragraph.runs):
            continue
        init = True
        for run in paragraph.runs:
            if init:
                init = False
                run.text = translated[p_index]
                p_index += 1
            else:
                run.text = ""

    document.save('output.docx')
    try:
        convert("output.docx", "output.pdf")
    except Exception:
        pass

    with open("output.docx", "rb") as f:
        docx_b64 = base64.b64encode(f.read()).decode('utf-8')
    with open("output.pdf", "rb") as f:
        pdf_b64 = base64.b64encode(f.read()).decode('utf-8')

    os.remove("temp.pdf")
    os.remove("input.docx")
    return [docx_b64, pdf_b64]


def get_all_translate_doc_list_api(user_id):
    res = kyzs_sql.mysql_exec(
        "SELECT id, name, status FROM translate_doc WHERE user_id=%s", (user_id,)
    )
    return {'translate_doc_list': res}


def get_translate_doc_detail_api(doc_id):
    # 数据库只存文件路径，读取时从磁盘读取并转 base64 返回前端
    res = kyzs_sql.mysql_exec(
        "SELECT output_pdf_base64, output_docx_base64 FROM translate_doc WHERE id=%s", (doc_id,)
    )[0]
    with open(res['output_pdf_base64'], 'rb') as f:
        pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
    with open(res['output_docx_base64'], 'rb') as f:
        docx_b64 = base64.b64encode(f.read()).decode('utf-8')
    return {'translate_doc_detail': {'output_pdf_base64': pdf_b64, 'output_docx_base64': docx_b64}}


def del_translate_doc_api(doc_id):
    res = kyzs_sql.mysql_exec(
        "SELECT output_pdf_base64, output_docx_base64 FROM translate_doc WHERE id=%s", (doc_id,)
    )[0]
    os.remove(res['output_pdf_base64'])
    os.remove(res['output_docx_base64'])
    kyzs_sql.mysql_exec("DELETE FROM translate_doc WHERE id=%s", (doc_id,))
    return 'ok'


def get_translate_word_by_content_api(content1: str):
    # LIKE 模糊查询同时搜索原文（content1）和译文（content2），LIMIT 50 防止返回过多
    res = kyzs_sql.mysql_exec(
        "SELECT * FROM translate_words WHERE content1 LIKE %s OR content2 LIKE %s LIMIT 50",
        (f'%{content1}%', f'%{content1}%')
    )
    if res:
        return {'code': 200, 'msg': 'success', 'data': res}
    return {'code': 404, 'msg': '词汇不存在', 'data': None}


def add_translate_word_api(ts_type: str, field_id: int, content1: str, content2: str, content3: str, from_source: str):
    try:
        kyzs_sql.mysql_exec(
            "INSERT INTO translate_words (ts_type, field_id, content1, content2, content3, `from`) VALUES (%s,%s,%s,%s,%s,%s)",
            (ts_type, field_id, content1, content2, content3, from_source)
        )
        new_id = kyzs_sql.mysql_exec("SELECT id FROM translate_words ORDER BY id DESC LIMIT 1")[0]['id']
        return {'code': 200, 'msg': 'success', 'data': {'id': new_id}}
    except Exception as e:
        return {'code': 500, 'msg': f'添加失败: {str(e)}', 'data': None}


def update_translate_word_api(word_id: int, ts_type: str = None, field_id: int = None,
                              content1: str = None, content2: str = None, content3: str = None,
                              from_source: str = None):
    """
    动态构建 UPDATE 语句，只更新传入的字段。
    字段名（列名）不能用 %s 参数化，但值全部通过 params 传递防止注入。
    """
    fields, params = [], []
    if ts_type is not None:
        fields.append("ts_type=%s"); params.append(ts_type)
    if field_id is not None:
        fields.append("field_id=%s"); params.append(field_id)
    if content1 is not None:
        fields.append("content1=%s"); params.append(content1)
    if content2 is not None:
        fields.append("content2=%s"); params.append(content2)
    if content3 is not None:
        fields.append("content3=%s"); params.append(content3)
    if from_source is not None:
        fields.append("`from`=%s"); params.append(from_source)

    if not fields:
        return {'code': 400, 'msg': '没有提供要更新的字段', 'data': None}

    params.append(word_id)
    try:
        kyzs_sql.mysql_exec(
            f"UPDATE translate_words SET {', '.join(fields)} WHERE id=%s",
            tuple(params)
        )
        return {'code': 200, 'msg': 'success', 'data': None}
    except Exception as e:
        return {'code': 500, 'msg': f'更新失败: {str(e)}', 'data': None}


def delete_translate_word_api(word_id: int):
    try:
        kyzs_sql.mysql_exec("DELETE FROM translate_words WHERE id=%s", (word_id,))
        return {'code': 200, 'msg': 'success', 'data': None}
    except Exception as e:
        return {'code': 500, 'msg': f'删除失败: {str(e)}', 'data': None}
