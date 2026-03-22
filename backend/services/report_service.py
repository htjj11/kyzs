from core.sqlLiteExec import sqlite_execute

#修改综述
def modify_review_new_api(review_id: int, review_body: str):
    sqlite_execute(
        "UPDATE review_records SET review_body=? WHERE id=?", (review_body, review_id)
    )
    return 1

#删除综述
def delete_summary(id: int):
    sqlite_execute("DELETE FROM review_records WHERE id=?", (id,))
    return 1

#传入综述id生成word文档
def generate_fuwenben_word_api(id: int):
    sql = "SELECT * FROM review_records WHERE id=?"
    print("\033[31m" + "获取综述内容用于生成word：" + sql + "\033[0m")
    data = sqlite_execute(sql, (id,))
    if not data:
        return {'code': 400, 'msg': '未找到该综述', 'data': None}

    html_text = data[0]['review_body']
    html_text = html_text.replace('\\n', '\n')

    from htmldocx import HtmlToDocx
    from docx import Document
    from docx.oxml.ns import qn
    import io

    doc = Document()
    parser = HtmlToDocx()
    parser.add_html_to_document(html_text, doc)

    # 统一设置文档所有的字体为宋体
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = u'宋体'
            if style._element.rPr is not None and style._element.rPr.rFonts is not None:
                style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = u'宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    buffer = io.BytesIO()
    doc.save(buffer)
    word_binary = buffer.getvalue()

    word_b64 = base64.b64encode(word_binary).decode('utf-8')
    return {'code': 200, 'msg': 'success', 'data': word_b64}
