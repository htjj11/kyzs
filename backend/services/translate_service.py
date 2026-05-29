import base64
import json
import re
import requests
import concurrent.futures
import time
import os

from core.sqlLiteExec import sqlite_execute
from services.第三方接口.大模型对话 import changcheng_ai_answer


# 翻译单个字符串
def translate_text_api(raw_text: str, translate_type: str, field_id: int):
    """
    翻译单个字符串，流程：
        1. 从 translate_words 词典表中查找原文里出现的专业术语
        2. 将匹配到的术语对照组拼入提示词，辅助大模型保持行业术语一致性
        3. 通过 Function Tool 强制模型返回标准 JSON {origin_text, translate_text}
        4. 解析失败时重试，直到成功 

    translate_type: 'en2zh' 英译中 / 'zh2en' 中译英
    field_id: 专业领域 id，用于过滤对应领域的词典
    """
    # ---------- 内部工具定义 ----------
    translate_tool = {
        "type": "function",
        "function": {
            "name": "output_translation",
            "description": "输出翻译结果，必须严格按照参数格式返回，不得添加任何额外文字",
            "parameters": {
                "type": "object",
                "properties": {
                    "origin_text": {
                        "type": "string",
                        "description": "原文，原样返回"
                    },
                    "translate_text": {
                        "type": "string",
                        "description": "译文"
                    }
                },
                "required": ["origin_text", "translate_text"]
            },
            "strict": False
        }
    }

    def call_translate(prompt: str) -> dict:
        """调用长城AI，通过提示词强制返回 JSON {origin_text, translate_text}。"""
        json_prompt = prompt + "\n\n请严格按照以下JSON格式返回，不要添加任何其他文字：\n{\"origin_text\": \"原文\", \"translate_text\": \"译文\"}"
        
        response_text = changcheng_ai_answer(json_prompt)
        
        # 尝试从返回文本中提取 JSON
        try:
            # 先尝试直接解析（strict=False 允许控制字符如 \t \n）
            result = json.loads(response_text, strict=False)
        except json.JSONDecodeError:
            # 如果直接解析失败，尝试提取 JSON 块
            match = re.search(r'\{[^{}]*"origin_text"[^{}]*"translate_text"[^{}]*\}', response_text, re.DOTALL)
            if match:
                result = json.loads(match.group(), strict=False)
            else:
                raise ValueError(f"无法从返回内容中提取JSON: {response_text}")
        
        print(f"长城AI翻译结果: {result}")
        return result
    # ---------- 主流程 ----------
    print('用户请求翻译:', raw_text[:40], '...')
    # en2zh：匹配 content1（英文）；zh2en：匹配 content2（中文）
    word_index = {'en2zh': 'content1', 'zh2en': 'content2'}
    # en2zh 查英文词条，zh2en 查中文词条
    ts_type_map = {'en2zh': 'en', 'zh2en': 'zh'}
    ts_type = ts_type_map[translate_type]

    words_dict = sqlite_execute(
        "SELECT * FROM translate_words WHERE ts_type=? AND field_id=?",
        (ts_type, field_id)
    ) or []

    col = word_index[translate_type]
    # 用正则全词匹配，避免"drill"匹配到"drilled"等子串；跳过字段为 None 的词条
    lower_paragraph = raw_text.lower()
    matched = [
        w for w in words_dict
        if w.get(col) is not None
        and re.search(r'\b' + re.escape(w[col].lower()) + r'\b', lower_paragraph)
    ]
    print('知识库命中字典个数:', len(matched))

    prompt = '以下是文中存在的词汇对照组，可用于参考：\n'
    for w in matched:
        prompt += f"{w[word_index[translate_type]]}：{w['content2']}，解释：{w['content3']}  "
    if translate_type == 'en2zh':
        prompt += f"\n请将以下英文文本翻译为中文：{raw_text}"
    else:
        prompt += f"\n请将以下中文文本翻译为英文：{raw_text}"

    while True:
        try:
            res_json = call_translate(prompt)
            translate_text = res_json['translate_text']
            print('大模型翻译成功：', translate_text)
            break
        except Exception as e:
            print('大模型返回的结果格式错误，重试中，错误信息:', e)

    return {'translate_result': translate_text, 'words_dict': matched}


# 翻译字符串列表
def translate_text_list_api(raw_text_list: list, translate_type: str, field_id: int):
    """
    分批并发翻译，规避 API 频控：
    1. 每 5 个段落为一组合并行翻译。
    2. 每执行完一批后等待 2 秒再开始下一批。
    """
    results = [None] * len(raw_text_list)
    batch_size = 5
    
    for start_idx in range(0, len(raw_text_list), batch_size):
        end_idx = min(start_idx + batch_size, len(raw_text_list))
        current_batch = raw_text_list[start_idx:end_idx]
        
        print(f"开始翻译批次: {start_idx} 到 {end_idx-1}")
        
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = {
                executor.submit(translate_text_api, text, translate_type, field_id): i
                for i, text in enumerate(current_batch)
            }
            for future in concurrent.futures.as_completed(futures):
                batch_rel_idx = futures[future]
                try:
                    results[start_idx + batch_rel_idx] = future.result()['translate_result']
                except Exception as e:
                    print(f"翻译第 {start_idx + batch_rel_idx} 段出错: {e}")
                    results[start_idx + batch_rel_idx] = "该段翻译失败"
        
        # 如果不是最后一批，则等待 2 秒
        if end_idx < len(raw_text_list):
            print("批次执行完成，等待 2 秒...")
            time.sleep(2)
            
    return results

# 创建翻译任务
def create_new_translate_doc_mission(base64_pdf_string, topic, user_id):
    """
    创建翻译任务并执行翻译，由路由层通过 BackgroundTasks 异步调用，不阻塞请求响应。
    执行完成后将翻译产物写入 ./file/translate_doc 目录，并将文件路径回写数据库（status=1 表示完成）。
    """
    sqlite_execute(
        "INSERT INTO translate_doc (name, status, raw_base64, user_id) VALUES (?, 0, ?, ?)",
        (topic, base64_pdf_string, user_id)
    )
    mission = sqlite_execute("SELECT id FROM translate_doc ORDER BY id DESC LIMIT 1")
    mission_id = mission[0]['id']

    output_docx_base64, output_pdf_base64 = translate_pdf(base64_pdf_string)
    file_name = time.strftime("%Y%m%d%H%M%S", time.localtime())

    # 翻译产物存本地 ./file/translate_doc 目录，数据库只存路径
    with open(f'./file_data/translate_doc/{file_name}.pdf', 'wb') as f:
        f.write(base64.b64decode(output_pdf_base64))
    with open(f'./file_data/translate_doc/{file_name}.docx', 'wb') as f:
        f.write(base64.b64decode(output_docx_base64))
 
    sqlite_execute(
        "UPDATE translate_doc SET status=1, output_pdf_base64=?, output_docx_base64=? WHERE id=?",
        (f'./file_data/translate_doc/{file_name}.pdf', f'./file_data/translate_doc/{file_name}.docx', mission_id)
    )
    print('翻译任务id成功:', mission_id)
    return 'ok'



# 翻译 PDF
def translate_pdf(base64_pdf_string):
    import tempfile
    import shutil
    import subprocess
    from pdf2docx import parse
    from docx import Document

    
    def get_all_paragraphs(document):
        """提取文档中所有段落（含表格内段落），返回 (paragraph, is_in_table) 列表"""
        result = []
        # 顶层段落
        for para in document.paragraphs:
            result.append(para)
        # 表格内段落
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        result.append(para)
        return result

    LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"
    temp_dir = tempfile.mkdtemp(prefix="translate_")

    try:
        # 1. 解码 PDF
        temp_pdf = os.path.join(temp_dir, "temp.pdf")
        with open(temp_pdf, "wb") as f:
            f.write(base64.b64decode(base64_pdf_string))

        # 2. PDF → docx
        input_docx = os.path.join(temp_dir, "input.docx")
        try:
            parse(temp_pdf, input_docx)
        except Exception as e:
            print(f"PDF解析出错: {e}")
        finally:
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)

        # 3. 读取所有段落（含表格内）
        document = Document(input_docx)
        all_paragraphs = get_all_paragraphs(document)

        print(f"顶层段落数: {len(document.paragraphs)}")
        print(f"总段落数（含表格）: {len(all_paragraphs)}")

        list_paragraphs = []   # 待翻译文本
        para_index_map = []    # 记录哪些段落需要翻译（用于回写）

        for para in all_paragraphs:
            text = para.text.strip()
            # 跳过空段落（无文字的纯图片段落也在此被跳过）
            if not text:
                continue
            # 合并 runs：只保留第一个有文字的 run，其余文字 run 清空
            first_text_run = None
            for run in para.runs:
                if run.text and run.text.strip() and first_text_run is None:
                    first_text_run = run
                elif run.text and run.text.strip():
                    run.text = ""
            if first_text_run is None:
                continue
            list_paragraphs.append(text)
            para_index_map.append(para)

        print(f"提取到 {len(list_paragraphs)} 个段落待翻译")

        if not list_paragraphs:
            # 兜底：直接返回原始 docx，pdf 为 None
            with open(input_docx, "rb") as f:
                docx_b64 = base64.b64encode(f.read()).decode('utf-8')
            return [docx_b64, None]

        # 4. 翻译
        translated = translate_text_list_api(list_paragraphs, 'en2zh', 1)

        # 5. 回写翻译结果
        for i, para in enumerate(para_index_map):
            first_text_written = False
            for run in para.runs:
                if run.text and run.text.strip() and not first_text_written:
                    run.text = translated[i]
                    first_text_written = True
                elif run.text and run.text.strip():
                    run.text = ""

        # 6. 保存翻译后 docx
        output_docx = os.path.join(temp_dir, "output.docx")
        document.save(output_docx)
        del document, translated, list_paragraphs

        if os.path.exists(input_docx):
            os.remove(input_docx)

        with open(output_docx, "rb") as f:
            docx_b64 = base64.b64encode(f.read()).decode('utf-8')

        # 7. LibreOffice 转 PDF
        pdf_b64 = None
        try:
            result = subprocess.run(
                [
                    LIBREOFFICE_PATH,
                    "--headless",
                    "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", temp_dir,
                    output_docx,
                ],
                timeout=60,
                capture_output=True,
                text=True,
            )
            # LibreOffice 输出文件名与输入同名，扩展名换成 .pdf
            output_pdf = os.path.join(temp_dir, "output.pdf")
            if result.returncode == 0 and os.path.exists(output_pdf):
                with open(output_pdf, "rb") as f:
                    pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
            else:
                print(f"LibreOffice转换失败，returncode={result.returncode}")
                print(f"stderr: {result.stderr}")
        except subprocess.TimeoutExpired:
            print("LibreOffice转换超时")
        except Exception as e:
            print(f"docx转PDF失败: {e}")

        return [docx_b64, pdf_b64]

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)




def add_translate_word_api(ts_type: str, field_id: int, content1: str, content2: str, content3: str, from_source: str):
    try:
        sqlite_execute(
            "INSERT INTO translate_words (ts_type, field_id, content1, content2, content3, `from`) VALUES (?,?,?,?,?,?)",
            (ts_type, field_id, content1, content2, content3, from_source)
        )
        new_id = sqlite_execute("SELECT id FROM translate_words ORDER BY id DESC LIMIT 1")[0]['id']
        return {'code': 200, 'msg': 'success', 'data': {'id': new_id}}
    except Exception as e:
        return {'code': 500, 'msg': f'添加失败: {str(e)}', 'data': None}


if __name__ == "__main__":
    # 测试某个文件翻译是否正常
    test_file_path = r"C:\Users\shuxi\Downloads\中英文文档 (1)\Enhancing Information Retrieval in the Drilling Domain_ Zero-Shot Learning ng.pdf"
    with open(test_file_path, 'rb') as f:
        base64_pdf_string = base64.b64encode(f.read()).decode('utf-8')
    translate_pdf(base64_pdf_string)
    print("翻译完成")
